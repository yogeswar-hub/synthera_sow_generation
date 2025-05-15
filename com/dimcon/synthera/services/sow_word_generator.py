import json
import logging
from datetime import datetime
from docx import Document

from com.dimcon.synthera.config.sow_template_path import mapping_template_path
from com.dimcon.synthera.utilities.log_handler import LoggerManager
from com.dimcon.synthera.utilities.document_styler import DocumentStyler
from com.dimcon.synthera.utilities.bedrock_rewriter import BedrockRewriter

from com.dimcon.synthera.resources.connect_aurora import get_engine
from com.dimcon.synthera.utilities.sessions_manager import DBSessionUtil
from com.dimcon.synthera.resources.organization_and_employees.employees import Employee
from com.dimcon.synthera.resources.statement_of_work.meeting_sow_json_store import MeetingSOWJsonStore

# Setup centralized logger
logger = LoggerManager.setup_logger(__name__, level=logging.DEBUG)

# Initialize database session utility
engine = get_engine()
db_util = DBSessionUtil(engine)


class SOWDocumentGenerator:
    def __init__(self, lead_id: int, mapping_template_path: str = mapping_template_path):
        self.lead_id = lead_id
        self.mapping_template_path = mapping_template_path
        self.template = {}
        self.answer_lookup = {}
        self.sub_answer_lookup = {}
        self.document = Document()
        self.rewriter = BedrockRewriter()

        # Apply global document styles
        DocumentStyler.apply_styles(self.document)

        # Fetch answers and initialize state
        self.sow_payload = self.get_raw_payload_from_db()

    def get_raw_payload_from_db(self):
        try:
            with db_util.session_scope() as session:
                row = (
                    session.query(MeetingSOWJsonStore)
                    .filter_by(lead_id=self.lead_id)
                    .order_by(MeetingSOWJsonStore.version_number.desc())
                    .first()
                )
                if not row:
                    raise ValueError(f"No SOW record found for lead_id: {self.lead_id}")

                logger.info(f"Fetched raw_payload for lead_id={self.lead_id}")

                version = row.version_number
                created_by = row.created_by
                created_at = row.created_at
                raw_payload = row.raw_payload

            self.version = version
            self.created_by = self.get_employee_name(created_by)
            self.created_date = created_at.strftime("%Y-%m-%d")
            return raw_payload

        except Exception as e:
            logger.error(f"Failed to fetch raw_payload: {e}", exc_info=True)
            raise

    def get_employee_name(self, emp_id):
        with db_util.session_scope() as session:
            emp = session.query(Employee).filter_by(emp_id=emp_id).first()
            return emp.employee_name if emp else "Unknown"

    def load_template(self):
        try:
            with open(self.mapping_template_path) as f:
                self.template = json.load(f)
            logger.info("Loaded mapping template successfully.")
        except Exception as e:
            logger.error(f"Failed to load mapping template: {e}", exc_info=True)
            raise

    def build_answer_lookup(self):
        try:
            for section in self.sow_payload.get("sections", []):
                for qa in section.get("questionsAndAnswers", []):
                    self.answer_lookup[qa["question_id"]] = qa["answers_found"]

                for subsection in section.get("subsections", []):
                    for qa in subsection.get("questionsAndAnswers", []):
                        self.sub_answer_lookup[qa["question_id"]] = qa["answers_found"]

            logger.info("Built answer and sub-answer lookup from raw payload.")
        except Exception as e:
            logger.error(f"Failed to build answer lookup: {e}", exc_info=True)
            raise

    def add_metadata(self):
        try:
            meta = self.template.get("documentMetadata", {})
            self.document.add_heading(meta.get("templateName", "SOW Document"), 0)
            self.document.add_paragraph(f"Version: {self.version}")
            self.document.add_paragraph(f"Created By: {self.created_by}")
            self.document.add_paragraph(f"Created Date: {self.created_date}")
            logger.info("Added metadata to document.")
        except Exception as e:
            logger.error(f"Failed to add metadata to document: {e}", exc_info=True)
            raise

    def add_table_of_contents(self):
        try:
            self.document.add_paragraph("Table of Contents", style="TOC Heading")

            for section in self.template.get("sections", []):
                section_title = section.get("sectionTitle", "Unnamed Section")
                section_order = section.get("sectionOrder")
                self.document.add_paragraph(f"{section_order}. {section_title}", style="TOC Entry")

                for subsection in section.get("subsections", []):
                    subsection_order = subsection.get("subsectionOrder")
                    subsection_title = subsection.get("subsectionTitle", "Unnamed Subsection")
                    self.document.add_paragraph(f"    {subsection_order} {subsection_title}", style="TOC SubEntry")

            self.document.add_paragraph().add_run().add_break()
            logger.info("Generated table of contents from mapping.")
        except Exception as e:
            logger.error("Failed to generate table of contents from mapping: %s", e, exc_info=True)
            raise

    def add_sections(self):
        try:
            for section in self.template.get("sections", []):
                order = section.get("sectionOrder")
                title = section.get("sectionTitle", f"Section {order}")
                self.document.add_heading(f"{order} {title}", level=1)

                for question in section.get("questions", []):
                    qid = question.get("questionId")
                    answers = self.answer_lookup.get(qid, [])
                    for answer in answers:
                        polished = self.rewriter.rewrite(answer)
                        self.document.add_paragraph(polished)

                for subsection in section.get("subsections", []):
                    sub_order = subsection.get("subsectionOrder")
                    sub_title = subsection.get("subsectionTitle", f"Subsection {sub_order}")
                    self.document.add_heading(f"{sub_order} {sub_title}", level=2)

                    for question in subsection.get("questions", []):
                        qid = question.get("questionId")
                        answers = self.sub_answer_lookup.get(qid, [])
                        for answer in answers:
                            polished = self.rewriter.rewrite(answer)
                            self.document.add_paragraph(polished)

            logger.info("Sections and subsections with answers added to document.")
        except Exception as e:
            logger.error(f"Failed to add sections and subsections: {e}", exc_info=True)
            raise

    def save(self, path: str):
        try:
            self.document.save(path)
            logger.info(f"Document saved to: {path}")
        except Exception as e:
            logger.error(f"Failed to save document to {path}: {e}", exc_info=True)
            raise
